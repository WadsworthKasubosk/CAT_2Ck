# -*- coding: utf-8 -*-
"""
将论文 md 转为格式化的 docx 文件（毕设论文风格）
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE_DIR, '论文_基于YOLO11与UNet的直肠肿瘤CT图像辅助诊断系统研究.md')
OUT_PATH = os.path.join(BASE_DIR, '论文_基于YOLO11与UNet的直肠肿瘤CT图像辅助诊断系统研究.docx')
FIG_DIR = os.path.join(BASE_DIR, '论文插图_png')

# 图号 → 文件名映射
FIGURE_MAP = {
    '2-1': '图2-1_UNet网络结构示意图.png',
    '2-2': '图2-2_YOLO11网络整体架构图.png',
    '2-3': '图2-3_YOLO11seg分割机制示意图.png',
    '3-1': '图3-1_系统总体架构图.png',
    '3-2': '图3-2_系统功能模块图.png',
    '3-3': '图3-3_数据库ER图.png',
    '4-1': '图4-1_数据预处理流程图.png',
    '4-2': '图4-2_YOLO标注格式示例.png',
    '4-3': '图4-3_推理流程对比图.png',
    '4-4': '图4-4_前端界面布局.png',
    '4-5': '图4-5_LLM辅助建议生成流程图.png',
    '5-1': '图5-1_训练损失曲线.png',
    '5-2': '图5-2_验证指标曲线.png',
    '5-3': '图5-3_验证集可视化.png',
    '5-4': '图5-4_模型对比.png',
    '5-5': '图5-5_系统主界面.png',
    '5-6': '图5-6_CT上传与分割结果.png',
    '5-7': '图5-7_影像组学特征值表格.png',
    '5-8': '图5-8_趋势分析图表.png',
    '5-9': '图5-9_LLM辅助建议.png',
    '5-10': '图5-10_患者管理界面.png',
}


def setup_styles(doc):
    """设置论文常用样式"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def add_title(doc, text):
    """论文大标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_heading1(doc, text):
    """一级标题（第X章）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(16)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_heading2(doc, text):
    """二级标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(14)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_heading3(doc, text):
    """三级标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(12)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_body(doc, text):
    """正文段落"""
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text.strip())
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_bold_body(doc, text):
    """加粗正文"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text.strip())
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_keywords(doc, text):
    """关键词行"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run_label = p.add_run('关键词：')
    run_label.font.name = '黑体'
    run_label.font.size = Pt(12)
    run_label.bold = True
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    # 提取关键词内容
    kw = text.replace('**关键词：**', '').replace('**关键词:**', '').strip()
    run_val = p.add_run(kw)
    run_val.font.name = '宋体'
    run_val.font.size = Pt(12)
    run_val._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_en_keywords(doc, text):
    """英文关键词"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run_label = p.add_run('Keywords: ')
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(12)
    run_label.bold = True
    kw = text.replace('**Keywords:**', '').replace('**Keywords：**', '').strip()
    run_val = p.add_run(kw)
    run_val.font.name = 'Times New Roman'
    run_val.font.size = Pt(12)


def add_code_block(doc, lines):
    """代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    code_text = '\n'.join(lines)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table(doc, header, rows):
    """插入表格"""
    ncols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h.strip())
        run.font.name = '黑体'
        run.font.size = Pt(10.5)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val.strip())
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表后空一行
    doc.add_paragraph()


def add_figure(doc, text):
    """插入图片 + 图题，找不到文件则留灰色占位"""
    # 从文本中提取图号，如 "图2-1" "图5-3"
    m = re.search(r'图(\d+-\d+)', text)
    fig_id = m.group(1) if m else None
    fig_file = FIGURE_MAP.get(fig_id) if fig_id else None
    fig_path = os.path.join(FIG_DIR, fig_file) if fig_file else None

    if fig_path and os.path.isfile(fig_path):
        # 插入图片
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run()
        run.add_picture(fig_path, width=Cm(14))

        # 图题
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        # 提取标题文字: 【图X-Y 标题】
        title_match = re.search(r'【(图\d+-\d+\s+.+?)】', text)
        cap_text = title_match.group(1) if title_match else f'图{fig_id}'
        run_c = cap.add_run(cap_text)
        run_c.font.name = '宋体'
        run_c.font.size = Pt(10.5)
        run_c._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    else:
        # 找不到文件，留占位
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        clean = text.strip().strip('*').strip()
        run = p.add_run(f'[{clean}]')
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def parse_table_block(lines, start):
    """从 md 行中解析表格，返回 (header, rows, end_index)"""
    # 找 header
    header_line = lines[start]
    header = [c.strip() for c in header_line.strip('|').split('|')]

    # 跳过分隔行
    idx = start + 1
    if idx < len(lines) and re.match(r'\|[\s\-:]+\|', lines[idx]):
        idx += 1

    rows = []
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        row = [c.strip() for c in lines[idx].strip('|').split('|')]
        rows.append(row)
        idx += 1

    return header, rows, idx


def build_docx():
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()
    setup_styles(doc)

    i = 0
    in_code = False
    code_lines = []
    first_h1 = True  # 大标题

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- 水平线（跳过）
        if stripped == '---':
            i += 1
            continue

        # 代码块
        if stripped.startswith('```'):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 论文大标题 (# 开头，只有一个)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            add_title(doc, stripped[2:].strip())
            i += 1
            continue

        # 章标题 (## 第X章 or ## 摘要 etc.)
        if stripped.startswith('## ') and not stripped.startswith('### '):
            text = stripped[3:].strip()
            if first_h1:
                # 摘要不分页
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(text)
                run.font.name = '黑体'
                run.font.size = Pt(16)
                run.bold = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                first_h1 = False
            else:
                add_heading1(doc, text)
            i += 1
            continue

        # 二级标题 ###
        if stripped.startswith('### ') and not stripped.startswith('#### '):
            add_heading2(doc, stripped[4:].strip())
            i += 1
            continue

        # 三级标题 ####
        if stripped.startswith('#### '):
            add_heading3(doc, stripped[5:].strip())
            i += 1
            continue

        # 关键词行
        if stripped.startswith('**关键词'):
            add_keywords(doc, stripped)
            i += 1
            continue
        if stripped.startswith('**Keywords'):
            add_en_keywords(doc, stripped)
            i += 1
            continue

        # 表格
        if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            header, rows, end_idx = parse_table_block(lines, i)
            add_table(doc, header, rows)
            i = end_idx
            continue

        # 图片标记 > **【图...
        if stripped.startswith('> **【图'):
            add_figure(doc, stripped)
            i += 1
            continue

        # 其他 > 行（图片说明的续行）
        if stripped.startswith('> '):
            inner = stripped[2:].strip()
            if inner.startswith('插入位置') or inner.startswith('内容') or inner.startswith('来源') or inner.startswith('建议'):
                # 图的补充说明，跳过
                i += 1
                continue
            else:
                add_body(doc, inner)
                i += 1
                continue

        # **表X-Y ...** 表标题
        if re.match(r'^\*\*表\d', stripped):
            text = stripped.strip('*').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(text)
            run.font.name = '黑体'
            run.font.size = Pt(10.5)
            run.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            i += 1
            continue

        # 加粗段落 **xxx**
        if stripped.startswith('**') and stripped.endswith('**'):
            add_bold_body(doc, stripped.strip('*').strip())
            i += 1
            continue

        # 普通正文
        # 清理行内 markdown
        text = stripped
        # 处理加粗
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        # 处理行内代码
        text = re.sub(r'`(.+?)`', r'\1', text)
        # 处理列表符号
        if text.startswith('- ') or text.startswith('* '):
            text = '· ' + text[2:]

        # 带编号的段落 (1) (2) 等
        if re.match(r'^\（?\d+[）\)]', text) or re.match(r'^\(\d+\)', text):
            add_body(doc, text)
            i += 1
            continue

        # 参考文献 [1] ...
        if re.match(r'^\[\d+\]', text):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(-0.74)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            i += 1
            continue

        add_body(doc, text)
        i += 1

    doc.save(OUT_PATH)
    print(f'已生成: {OUT_PATH}')


if __name__ == '__main__':
    build_docx()
